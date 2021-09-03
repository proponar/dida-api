import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

import re
import io
import sys
import shutil

if len(sys.argv) > 1:
    data = json.load(open(sys.argv[1])) # '/home/martin/Projects/word-py/ex.json'))
else:
    data = json.load(sys.stdin)
    #with open('/tmp/dump.js', 'w') as out:
    #    json.dump(data, out) # debug dump

document = Document()

# change page orientation
current_section = document.sections[-1]
new_width, new_height = current_section.page_height, current_section.page_width
current_section.orientation = WD_ORIENT.LANDSCAPE
current_section.page_width = new_width
current_section.page_height = new_height

# debug dump
#p = document.add_paragraph(json.dumps(data))

# set styles
styles = document.styles
styles['Heading 1'].font.color.rgb = RGBColor(0xBF, 0x00, 0x41)
styles['Heading 1'].font.underline = False
styles['Heading 1'].font.size = Pt(28)

styles['Heading 2'].font.underline = True
styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
styles['Heading 2'].font.size = Pt(16)

styles['Heading 3'].font.underline = False
styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
styles['Heading 3'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles['Heading 3'].font.size = Pt(14)

same_heslo = ''
same_meaning = ''
same_urceni = ''

def format_urceni(u):
    if u == None or u["pad"] == None:
        return ""

    md = re.match("(\d)(\w)", u["pad"])
    if not md:
        return ""

    num_str = "sg." if md.groups()[1] == "s" else "pl."

    urceni_string = f'{md.groups()[0]} {num_str}'

    if u["rod"] != None and u["rod"] != " ":
        urceni_string = f'{urceni_string} {u["rod"]}'

    return urceni_string

def format_lokalizace(e):
    if 'lokalizace_format' in e and e['lokalizace_format'] != '':
        return e['lokalizace_format']
    if 'lokalizace_text' in e and e['lokalizace_text'] != '':
        return e['lokalizace_text']
    return ''

def format_zdroj(z):
    l = []
    # autor - název1 - název2 - rok
    if 'autor' in z and z['autor'] != None:
        l.append(z['autor'])
    if 'name' in z and z['name'] != None:
        l.append(z['name'])
    if 'nazev2' in z and z['nazev2'] != None:
        l.append(z['nazev2'])
    if 'rok' in z and z['rok'] != None:
        l.append(str(z['rok']))
    return " - ".join(l)

def format_cislo_vyznamu(v):
    if v != None:
        return v
    else:
        return ''

for e in data:
    if same_heslo != e['heslo']:
        heslo_text = f'{e["heslo"]} {e["entry_full"]["rod"]}.'
        document.add_heading(heslo_text, 1)
        same_heslo = e['heslo']
        same_meaning = ''
        same_urceni = ''

    if same_meaning != e['vyznam']:
        #document.add_heading(vyznam_text, 2)

        p = document.add_paragraph(f'{format_cislo_vyznamu(e["vyznam_full"]["cislo"])}. ', style='Heading 2')
        if None != e["vyznam_full"]["kvalifikator"]:
            p.add_run(e["vyznam_full"]["kvalifikator"]).bold = False
        p.add_run(e["vyznam"])

        same_meaning = e['vyznam']
        same_urceni = ''

    if same_urceni != e['urceni_sort']:
        if len(e['urceni_full']) > 0 :
            urceni_string = format_urceni(e['urceni_full'][0])
        else:
            urceni_string = ''
        document.add_heading(urceni_string, 3)
        same_urceni = e['urceni_sort']

    p = document.add_paragraph('', style='List Bullet')
    # splitnout exemplifikaci podle vyrazu, vyrazy tucne
    parts = re.split("(\{.*?\})", e['exemplifikace'])
    for part in parts:
        if part.startswith('{'):
            p.add_run(part).bold = True
        else:
            p.add_run(part).italic = True

    p.add_run('; ')

    # lokalizace cervene
    lokalizace = format_lokalizace(e)
    if lokalizace != '':
        run = p.add_run(lokalizace)
        font = run.font
        font.size = Pt(10)
        font.color.rgb = RGBColor(0xE9, 0x24, 0x42)
        p.add_run('; ')

    # zdroj modre
    if 'zdroj_name' in e:
        run = p.add_run(format_zdroj(e["source_full"]))
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

# output the document to STDOUT
target_stream = io.BytesIO()
document.save(target_stream)
target_stream.seek(0)
sys.stdout.buffer.write(target_stream.read())
